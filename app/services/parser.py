"""
Document parsing service.

Each parser function takes raw bytes and returns (text, metadata, warnings).
Routes call `parse_file()`, which dispatches to the right handler based on
filename extension.

Design notes:
- We never accept arbitrary file types — only the kinds in FileKind enum.
- Parsers are deliberately separate functions (not classes) for testability.
- We surface non-fatal issues as warnings, not exceptions, per the tiered policy.
- Hard failures (corrupted file, password-protected, fully scanned) raise.
"""

from __future__ import annotations

import io
from dataclasses import dataclass
from pathlib import Path

import pdfplumber
from docx import Document as DocxDocument

from app.schemas import FileKind, ParseMetadata, ParseWarning

# ─── thresholds (tunable) ──────────────────────────────────────────────────

# Below this much extractable text, we assume it's a scanned PDF. Picked
# empirically: a typical 1-page resume has 1500-3500 chars; 30 chars is "this
# is almost certainly an image with no text layer."
MIN_TOTAL_CHARS = 30

# Per-page density threshold for a "low text density" warning. Resumes have
# 800-2000 chars/page typically; below 80 chars/page suggests something's off
# (scanned but with a fluke OCR layer, or a heavily image-based design).
MIN_CHARS_PER_PAGE = 80

# Horizontal-gap threshold for word boundaries during PDF text extraction.
# pdfplumber's default is 3, which fails on tightly-kerned PDFs (LaTeX
# templates, Canva exports, designer resumes) — words run together with no
# spaces. 1.5 handles tight typography while still recognizing word breaks
# in normally-spaced PDFs. Lower = more aggressive about inserting spaces.
PDF_X_TOLERANCE = 1.5


# ─── exceptions ────────────────────────────────────────────────────────────


class ParserError(Exception):
    """Base class for parser failures. Routes catch this and return 4xx."""

    def __init__(self, code: str, message: str, detail: str | None = None):
        self.code = code
        self.message = message
        self.detail = detail
        super().__init__(message)


class UnsupportedFileType(ParserError):
    pass


class ScannedPDFError(ParserError):
    pass


class EncryptedPDFError(ParserError):
    pass


class CorruptedFileError(ParserError):
    pass


# ─── result container ──────────────────────────────────────────────────────


@dataclass
class ParseResult:
    """Internal struct returned by parsers. Routes convert to ParseResponse."""

    text: str
    metadata: ParseMetadata
    warnings: list[ParseWarning]


# ─── kind detection ────────────────────────────────────────────────────────


def detect_kind(filename: str) -> FileKind:
    """Map a filename to a FileKind. Raises UnsupportedFileType on miss."""
    ext = Path(filename).suffix.lower().lstrip(".")
    try:
        return FileKind(ext)
    except ValueError as e:
        raise UnsupportedFileType(
            code="unsupported_file_type",
            message=f"File extension '.{ext}' is not supported.",
            detail=f"Supported types: {', '.join(k.value for k in FileKind)}",
        ) from e


# ─── parsers per file type ─────────────────────────────────────────────────


def parse_pdf(data: bytes, filename: str) -> ParseResult:
    """Extract text from a PDF using pdfplumber."""
    warnings: list[ParseWarning] = []
    text_parts: list[str] = []
    pages_with_text = 0

    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                page_text = page.extract_text(x_tolerance=PDF_X_TOLERANCE) or ""
                if page_text.strip():
                    pages_with_text += 1
                text_parts.append(page_text)
    except Exception as e:
        # pdfplumber raises various opaque errors for encryption/corruption.
        # Sniff the message to give a useful response.
        msg = str(e).lower()
        if "password" in msg or "encrypted" in msg:
            raise EncryptedPDFError(
                code="encrypted_pdf",
                message="This PDF is password-protected.",
                detail="Remove the password (Preview → Export, or any PDF tool) and re-upload.",
            ) from e
        raise CorruptedFileError(
            code="corrupted_pdf",
            message="This PDF could not be parsed.",
            detail=str(e),
        ) from e

    text = "\n\n".join(text_parts).strip()
    char_count = len(text)

    # Scanned-PDF detection (tiered policy)
    if char_count < MIN_TOTAL_CHARS:
        raise ScannedPDFError(
            code="scanned_pdf",
            message="This PDF appears to be scanned — no text layer was found.",
            detail=(
                f"Extracted only {char_count} characters across {page_count} page(s). "
                "Run OCR first (Adobe Acrobat, Preview, or an online tool), "
                "or paste the text directly."
            ),
        )

    # Low-density warning
    chars_per_page = char_count / max(page_count, 1)
    if chars_per_page < MIN_CHARS_PER_PAGE:
        warnings.append(ParseWarning.LOW_TEXT_DENSITY)

    # Partial-extraction warning (some pages came back empty)
    if pages_with_text < page_count:
        warnings.append(ParseWarning.PARTIAL_EXTRACTION)

    metadata = ParseMetadata(
        kind=FileKind.PDF,
        filename=filename,
        size_bytes=len(data),
        page_count=page_count,
        char_count=char_count,
    )
    return ParseResult(text=text, metadata=metadata, warnings=warnings)


def parse_docx(data: bytes, filename: str) -> ParseResult:
    """Extract text from a .docx using python-docx."""
    try:
        doc = DocxDocument(io.BytesIO(data))
    except Exception as e:
        raise CorruptedFileError(
            code="corrupted_docx",
            message="This .docx could not be parsed.",
            detail=str(e),
        ) from e

    # Paragraphs cover the body; tables are pulled separately.
    parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    text = "\n".join(parts).strip()
    metadata = ParseMetadata(
        kind=FileKind.DOCX,
        filename=filename,
        size_bytes=len(data),
        page_count=None,  # docx has no fixed pagination
        char_count=len(text),
    )
    return ParseResult(text=text, metadata=metadata, warnings=[])


def parse_text(data: bytes, filename: str, kind: FileKind) -> ParseResult:
    """Plain text passthrough — works for .txt, .md, .tex."""
    try:
        text = data.decode("utf-8").strip()
    except UnicodeDecodeError:
        # Try the next-most-common encoding before giving up
        try:
            text = data.decode("latin-1").strip()
        except UnicodeDecodeError as e:
            raise CorruptedFileError(
                code="encoding_error",
                message="Could not decode file as UTF-8 or Latin-1.",
                detail=str(e),
            ) from e

    metadata = ParseMetadata(
        kind=kind,
        filename=filename,
        size_bytes=len(data),
        page_count=None,
        char_count=len(text),
    )
    return ParseResult(text=text, metadata=metadata, warnings=[])


# ─── public entry point ────────────────────────────────────────────────────


def parse_file(data: bytes, filename: str) -> ParseResult:
    """
    Dispatch to the right parser. Single entry point used by the route.

    Raises ParserError subclasses on failure. Returns ParseResult on success
    (which may include non-fatal warnings).
    """
    kind = detect_kind(filename)

    if kind == FileKind.PDF:
        return parse_pdf(data, filename)
    if kind == FileKind.DOCX:
        return parse_docx(data, filename)
    # txt, md, tex all go through the same plain-text path
    return parse_text(data, filename, kind)
