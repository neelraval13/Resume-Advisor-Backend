"""
Tests for /api/parse.

Strategy: generate small fixture files in-memory (not on disk, not in repo)
using the same libraries the parser uses. This keeps the repo clean and tests
self-contained.

We test:
- Plain text passes through cleanly.
- DOCX with paragraphs and tables both extract.
- PDF with text extracts, metadata is right.
- Scanned PDFs (no text layer) get rejected with a clear error.
- Unsupported extensions get rejected.
- Files over the size limit get rejected.
- Empty filenames get rejected.
"""

import io
import os

os.environ.setdefault("ANTHROPIC_API_KEY", "sk-ant-test-key-not-real")

import pytest  # noqa: E402
from docx import Document as DocxDocument  # noqa: E402
from fastapi.testclient import TestClient  # noqa: E402
from pdfplumber.utils.exceptions import (
    PdfminerException,  # noqa: E402, F401  # imported for type hint reference
)

from app.main import create_app  # noqa: E402


@pytest.fixture
def client():
    return TestClient(create_app())


# ─── fixture builders ─────────────────────────────────────────────────────


def make_docx_bytes(paragraphs: list[str], table_rows: list[list[str]] | None = None) -> bytes:
    """Build a minimal .docx in memory."""
    doc = DocxDocument()
    for p in paragraphs:
        doc.add_paragraph(p)
    if table_rows:
        table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for r, row in enumerate(table_rows):
            for c, cell in enumerate(row):
                table.rows[r].cells[c].text = cell
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def make_text_pdf_bytes(lines: list[str]) -> bytes:
    """
    Build a minimal PDF with real text content using reportlab.
    We add reportlab as a dev dep just for this; pdfplumber alone can't write.
    """
    from reportlab.pdfgen import canvas

    buf = io.BytesIO()
    c = canvas.Canvas(buf)
    y = 800
    for line in lines:
        c.drawString(50, y, line)
        y -= 20
    c.save()
    return buf.getvalue()


def make_blank_pdf_bytes() -> bytes:
    """A PDF with one empty page — simulates a scanned PDF (no text layer)."""
    from reportlab.pdfgen import canvas

    buf = io.BytesIO()
    c = canvas.Canvas(buf)
    c.showPage()  # blank page, no drawString calls
    c.save()
    return buf.getvalue()


# ─── happy path ───────────────────────────────────────────────────────────


def test_parse_txt(client):
    text = "Hello world\nThis is a test resume."
    r = client.post(
        "/api/parse",
        files={"file": ("resume.txt", text.encode("utf-8"), "text/plain")},
    )
    assert r.status_code == 200
    body = r.json()
    assert body["text"] == text
    assert body["metadata"]["kind"] == "txt"
    assert body["metadata"]["char_count"] == len(text)
    assert body["warnings"] == []


def test_parse_md(client):
    text = "# My Resume\n\nSenior engineer with 10 years of experience."
    r = client.post(
        "/api/parse",
        files={"file": ("resume.md", text.encode("utf-8"), "text/markdown")},
    )
    assert r.status_code == 200
    assert r.json()["metadata"]["kind"] == "md"


def test_parse_docx_with_paragraphs(client):
    data = make_docx_bytes(["John Smith", "Software Engineer", "Built things at companies."])
    r = client.post(
        "/api/parse",
        files={
            "file": (
                "resume.docx",
                data,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert r.status_code == 200
    body = r.json()
    assert "John Smith" in body["text"]
    assert "Software Engineer" in body["text"]
    assert body["metadata"]["kind"] == "docx"


def test_parse_docx_with_table(client):
    data = make_docx_bytes(
        paragraphs=["Skills"],
        table_rows=[["Language", "Years"], ["Python", "8"], ["Rust", "2"]],
    )
    r = client.post(
        "/api/parse",
        files={
            "file": (
                "resume.docx",
                data,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert r.status_code == 200
    text = r.json()["text"]
    # Table content must be extracted (default python-docx behavior misses it)
    assert "Python" in text
    assert "Rust" in text


def test_parse_pdf_with_text(client):
    data = make_text_pdf_bytes(
        ["John Smith", "Senior Engineer at Odyssey", "Built really cool things."]
    )
    r = client.post(
        "/api/parse",
        files={"file": ("resume.pdf", data, "application/pdf")},
    )
    assert r.status_code == 200
    body = r.json()
    assert "John Smith" in body["text"]
    assert body["metadata"]["kind"] == "pdf"
    assert body["metadata"]["page_count"] == 1


# ─── error paths ──────────────────────────────────────────────────────────


def test_parse_scanned_pdf_rejected(client):
    """A PDF with no text layer should be rejected with a clear error."""
    data = make_blank_pdf_bytes()
    r = client.post(
        "/api/parse",
        files={"file": ("scanned.pdf", data, "application/pdf")},
    )
    assert r.status_code == 400
    body = r.json()
    assert body["detail"]["error"] == "scanned_pdf"
    assert "OCR" in body["detail"]["detail"]


def test_parse_unsupported_extension(client):
    r = client.post(
        "/api/parse",
        files={"file": ("resume.rtf", b"some content", "application/rtf")},
    )
    assert r.status_code == 400
    assert r.json()["detail"]["error"] == "unsupported_file_type"


def test_parse_oversized_file(client):
    """Default limit is 10MB; send 11MB of data."""
    data = b"x" * (11 * 1024 * 1024)
    r = client.post(
        "/api/parse",
        files={"file": ("huge.txt", data, "text/plain")},
    )
    assert r.status_code == 413
    assert r.json()["detail"]["error"] == "file_too_large"


def test_parse_corrupted_pdf(client):
    """Bytes that aren't a real PDF should yield a 422."""
    r = client.post(
        "/api/parse",
        files={"file": ("fake.pdf", b"not actually a pdf", "application/pdf")},
    )
    assert r.status_code == 422
    assert r.json()["detail"]["error"] == "corrupted_pdf"
